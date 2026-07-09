"""On-the-fly merge of a base PDF + a signature-certificate annex.

Why an annex instead of modifying the original?
================================================
The signed payload of every ЭЦП is the **DOCX bytes** (SHA-256 hash bound into
the CMS structure). Touching the DOCX after signing invalidates every existing
``Signature`` row. But the PDF served as the human-readable view is NOT signed
— it is a one-way derivative of the signed DOCX. Modifying the PDF (or
appending pages to it) carries zero legal risk because nobody verifies the
PDF cryptographically: verification happens against the original DOCX bytes,
which we keep untouched.

This module produces, on demand, the PDF a user actually downloads:

  [ base PDF ──── original contract / consent / договор практики ]
  [ +annex ──── "Сертификат подписания": signers, masked IIN, time,
                verification level, QR code, public verify URL ]

Strategy
========
- Generate the signature-certificate DOCX, convert to PDF via the existing
  ``_convert_to_pdf`` (LibreOffice). Cache the result under the same archive
  subtree so a re-download doesn't pay the LibreOffice cost twice.
- Merge using ``pypdf`` — fast, no external process, preserves font subsets.
- Yield the merged file path. Caller streams it via ``send_file``.

Failure paths
=============
- If the base PDF doesn't exist (DOCX-only generation, LibreOffice missing in
  prod): return ``None`` and let the caller fall back to the raw DOCX or
  show an explicit error.
- If the certificate generator throws: log it and serve the base PDF as-is —
  the user still gets their original file; a missing visualization annex is
  not a reason to deny the download.
- If pypdf merge throws: same fallback.

NOTE on caching
================
The annex includes "Сформирован сертификат: dd.mm.YYYY HH:MM UTC" which means
the merged PDF will differ on every regeneration. We DO cache the certificate
PDF for a short window so two clicks within the same second produce identical
files; subsequent generations refresh the timestamp. This is a UX nicety, not
a correctness requirement — the verify URL is what matters legally.
"""
from __future__ import annotations

import hashlib
import io
import tempfile
from datetime import datetime, timedelta
from pathlib import Path
from typing import Callable

from flask import current_app
from pypdf import PdfReader, PdfWriter


_CERTIFICATE_CACHE_TTL = timedelta(minutes=5)

# Salt module-cache — populated lazily on first _renderer_version() call.
_RENDERER_VERSION_CACHED: str | None = None


def _cache_dir() -> Path:
    """Per-request short-lived merge cache (cleared on container restart)."""
    base = Path(tempfile.gettempdir()) / "ccu_pdf_merge_cache"
    base.mkdir(parents=True, exist_ok=True)
    return base


def _renderer_version() -> str:
    """Salt the cache key with a hash of the certificate/report renderers'
    source content. ANY change to how signatures are rendered — a label
    rename, a new field, a wording tweak — busts every cached merged PDF
    on the next request. Prevents "I deployed a fix but the old PDF is
    still being served for 5 minutes" surprises.

    Cached on the module: computed once per process (import cost only).
    """
    global _RENDERER_VERSION_CACHED
    if _RENDERER_VERSION_CACHED is not None:
        return _RENDERER_VERSION_CACHED
    h = hashlib.sha256()
    for rel in (
        # Certificate/report renderers whose output is embedded in the merge.
        "enrollment_certificate.py",   # enrollment (contract + consent) annex
        "signature_report.py",         # practicum three-party annex
        # LMS annex builder lives inline inside pdf_visualization.py, so its
        # own module content salts the hash below.
        "pdf_visualization.py",
    ):
        try:
            p = Path(__file__).with_name(rel)
            if p.is_file():
                h.update(p.read_bytes())
        except OSError:
            pass
    _RENDERER_VERSION_CACHED = h.hexdigest()[:8]
    return _RENDERER_VERSION_CACHED


def _signature_fingerprint(signatures) -> str:
    """Hash that changes whenever the signature set on the entity changes.

    Used as part of the cache key so a freshly-added signature invalidates a
    cached merged PDF without having to track creation timestamps explicitly.
    Also salted with the RENDERER version so a code change to how signatures
    are visualized busts every prior cache entry on the next request.
    """
    h = hashlib.sha256()
    h.update(_renderer_version().encode())
    for s in sorted(signatures or [], key=lambda x: (x.id or 0)):
        h.update(str(s.id).encode())
        h.update((s.signed_payload_sha256 or "").encode())
        h.update((s.created_at.isoformat() if s.created_at else "").encode())
    return h.hexdigest()[:16]


def merge_pdf_with_certificate(
    *,
    base_pdf_path: Path,
    cache_key: str,
    certificate_builder: Callable[[Path], Path | None],
) -> Path | None:
    """Merge ``base_pdf_path`` with the certificate produced by ``certificate_builder``.

    Parameters
    ----------
    base_pdf_path:
        Absolute path to the original (signed-derived) PDF — the contract,
        consent, practicum, or LMS document.
    cache_key:
        Stable identifier for this (entity, signature-set) tuple. The merged
        output is keyed by ``cache_key`` and reused for 5 minutes.
    certificate_builder:
        Callback that, given an output directory, returns the absolute path of
        the certificate PDF to APPEND. May return ``None`` if the certificate
        couldn't be produced (e.g. LibreOffice unavailable) — in that case we
        return the base PDF path unchanged.

    Returns
    -------
    Absolute path to the merged PDF, or the base path when no merge happened,
    or ``None`` if the base PDF doesn't exist on disk.
    """
    if not base_pdf_path.is_file():
        return None

    cache = _cache_dir()
    out_path = cache / f"{cache_key}.pdf"
    now = datetime.utcnow()
    if out_path.is_file():
        age = now - datetime.utcfromtimestamp(out_path.stat().st_mtime)
        if age < _CERTIFICATE_CACHE_TTL:
            return out_path

    try:
        cert_path = certificate_builder(cache)
    except Exception:  # noqa: BLE001
        current_app.logger.exception(
            "Certificate builder threw — serving base PDF as-is for %s",
            base_pdf_path.name,
        )
        return base_pdf_path

    if not cert_path or not Path(cert_path).is_file():
        # LibreOffice may be missing; fall back to the base file.
        current_app.logger.warning(
            "Certificate PDF not produced — serving base PDF as-is for %s",
            base_pdf_path.name,
        )
        return base_pdf_path

    try:
        writer = PdfWriter()
        for src in (base_pdf_path, Path(cert_path)):
            with open(src, "rb") as fh:
                reader = PdfReader(fh)
                for page in reader.pages:
                    writer.add_page(page)
        # Write to a temp + rename for atomicity (parallel requests for the
        # same cache_key won't see a half-written file).
        tmp = out_path.with_suffix(".pdf.tmp")
        with open(tmp, "wb") as fh:
            writer.write(fh)
        tmp.replace(out_path)
        return out_path
    except Exception:
        current_app.logger.exception(
            "PDF merge failed — serving base PDF as-is for %s", base_pdf_path.name,
        )
        return base_pdf_path


# ────────────────────────────────────────────────────────────────────────────
# Convenience wrappers — one per aggregate. Each computes its own cache key
# and certificate builder, so the API layer just calls one of these.
# ────────────────────────────────────────────────────────────────────────────

def merge_enrollment_pdf(e, base_pdf_path: Path, *, public_base: str | None = None) -> Path | None:
    """Merge an enrollment contract/consent PDF with its signature certificate."""
    from .enrollment_certificate import generate_signature_certificate
    if not (e.signatures or []):
        # No signatures → no annex to append. Serve the base file.
        return base_pdf_path if base_pdf_path.is_file() else None

    fingerprint = _signature_fingerprint(e.signatures)
    cache_key = f"enroll_{e.id}_{base_pdf_path.stem}_{fingerprint}"

    def _builder(_out_dir: Path) -> Path | None:
        # Reuse the enrollment_certificate generator — it writes under the
        # archive tree, so we just hand off its PDF path. The DOCX is also
        # written but we only need the PDF for the merge.
        try:
            _docx_path, pdf_path = generate_signature_certificate(e, public_base=public_base)
        except Exception:
            current_app.logger.exception("enrollment certificate gen failed")
            return None
        return pdf_path

    return merge_pdf_with_certificate(
        base_pdf_path=base_pdf_path,
        cache_key=cache_key,
        certificate_builder=_builder,
    )


def merge_practicum_pdf(contract, base_pdf_path: Path, *, public_base: str | None = None) -> Path | None:
    """Merge a three-party practicum contract PDF with its signature report.

    Reuses ``services.signature_report.build_signature_report`` which already
    produces a brand-styled DOCX listing each ЭЦП signer (college/partner/
    student). LibreOffice converts it to PDF; we append after the original
    contract pages.
    """
    from .signature_report import build_signature_report
    if not (contract.signatures or []):
        return base_pdf_path if base_pdf_path.is_file() else None

    fingerprint = _signature_fingerprint(contract.signatures)
    cache_key = f"practicum_{contract.id}_{base_pdf_path.stem}_{fingerprint}"

    def _builder(_out_dir: Path) -> Path | None:
        try:
            report_docx = build_signature_report(contract)
        except Exception:
            current_app.logger.exception("practicum signature report gen failed")
            return None
        report_pdf = report_docx.with_suffix(".pdf")
        if not report_pdf.is_file():
            from .document_generator import _convert_to_pdf
            produced = _convert_to_pdf(report_docx)
            if not produced or not Path(produced).is_file():
                return None
            report_pdf = produced
        return report_pdf

    return merge_pdf_with_certificate(
        base_pdf_path=base_pdf_path,
        cache_key=cache_key,
        certificate_builder=_builder,
    )


def merge_lms_pdf(lms, base_pdf_path: Path, *, public_base: str | None = None) -> Path | None:
    """Merge an LMS contract PDF with a single-signer signature page.

    The LMS aggregate is single-document + single-signer (student or parent)
    so we inline a small builder rather than spinning up a dedicated module.
    """
    from .enrollment_certificate import (
        _BRAND_CORAL, _BRAND_CHARCOAL, _BRAND_INK,
        _h, _p, _kv, _mask_iin, _serial_tail, _LEVEL_RU, _fmt_date,
    )
    from .document_generator import _convert_to_pdf
    from .qr import generate_qr_png, public_verify_url
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Cm, Pt
    from ..utils.time import utc_now

    if not (lms.signatures or []):
        return base_pdf_path if base_pdf_path.is_file() else None

    fingerprint = _signature_fingerprint(lms.signatures)
    cache_key = f"lms_{lms.id}_{base_pdf_path.stem}_{fingerprint}"

    def _builder(out_dir: Path) -> Path | None:
        try:
            verify_url = public_verify_url(lms.verify_code, base_url=public_base) if lms.verify_code else None
            doc = Document()
            section = doc.sections[0]
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

            _h(doc, "Сертификат подписания · LMS-договор", size=22)
            _h(doc, f"Договор № {lms.number or '—'} от {_fmt_date(lms.contract_date)}",
               size=13, color=_BRAND_CHARCOAL, bold=True)
            _p(doc, "CCU PRACTICUM · College of Caspian University",
               size=9, color=_BRAND_CHARCOAL, italic=True)
            doc.add_paragraph()

            _h(doc, "Сведения о договоре", size=14, color=_BRAND_CHARCOAL)
            _kv(doc, "Номер договора", lms.number or "—", bold_value=True)
            _kv(doc, "Дата договора", _fmt_date(lms.contract_date))
            _kv(doc, "Студент (грантник)", lms.applicant_full_name or "—", bold_value=True)
            _kv(doc, "ИИН (скрыт)", _mask_iin(lms.applicant_iin))
            _kv(doc, "Специальность", lms.specialty or "—")
            if lms.qualification:
                _kv(doc, "Квалификация", lms.qualification)
            _kv(doc, "Финансирование", lms.funding_source or "—")
            if lms.grant_order_number:
                _kv(doc, "Приказ о зачислении", lms.grant_order_number)
            _kv(doc, "Сформирован сертификат", utc_now().strftime("%d.%m.%Y %H:%M UTC"))
            doc.add_paragraph()

            _h(doc, "Список подписей", size=14, color=_BRAND_CHARCOAL)
            for idx, s in enumerate(sorted(lms.signatures, key=lambda x: x.created_at or 0)):
                _h(doc, f"Подпись №{idx + 1}", size=12, color=_BRAND_INK)
                _kv(doc, "  Подписант", s.signer_full_name or "—", bold_value=True)
                _kv(doc, "  ИИН (скрыт)", _mask_iin(s.signer_iin_or_bin))
                _kv(doc, "  Серийный № сертификата", _serial_tail(s.signer_serial))
                _kv(doc, "  Время подписания",
                    s.created_at.strftime("%d.%m.%Y %H:%M:%S UTC") if s.created_at else "—")
                _kv(doc, "  Уровень верификации",
                    _LEVEL_RU.get(s.verification_level, s.verification_level or "—"))
                if s.signed_payload_sha256:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    rl = p.add_run("  SHA-256: ")
                    rl.bold = True
                    rl.font.color.rgb = _BRAND_CHARCOAL
                    rl.font.size = Pt(9)
                    rl.font.name = "Inter"
                    rv = p.add_run(s.signed_payload_sha256)
                    rv.font.size = Pt(9)
                    rv.font.color.rgb = _BRAND_INK
                    rv.font.name = "Consolas"
                doc.add_paragraph()

            if verify_url:
                _h(doc, "Проверка подлинности", size=14, color=_BRAND_CHARCOAL)
                tbl = doc.add_table(rows=1, cols=2)
                tbl.autofit = False
                qr_cell, info_cell = tbl.rows[0].cells
                qr_cell.width = Cm(4.5)
                info_cell.width = Cm(12)
                png = generate_qr_png(verify_url, box_size=10, border=2)
                qp = qr_cell.paragraphs[0]
                qp.add_run().add_picture(io.BytesIO(png), width=Cm(3.8))
                info_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                ip = info_cell.paragraphs[0]
                r = ip.add_run("Отсканируйте QR-код или откройте ссылку:")
                r.font.size = Pt(10); r.font.color.rgb = _BRAND_CHARCOAL
                p = info_cell.add_paragraph()
                r = p.add_run(verify_url)
                r.font.size = Pt(11); r.font.color.rgb = _BRAND_CORAL; r.bold = True
                p = info_cell.add_paragraph()
                r = p.add_run(f"Код документа: {lms.verify_code}")
                r.font.size = Pt(9); r.font.color.rgb = _BRAND_CHARCOAL; r.italic = True

            tmp_docx = out_dir / f"_lms_cert_{lms.id}.docx"
            doc.save(str(tmp_docx))
            pdf_out = _convert_to_pdf(tmp_docx)
            return pdf_out if pdf_out and Path(pdf_out).is_file() else None
        except Exception:
            current_app.logger.exception("LMS cert gen failed")
            return None

    return merge_pdf_with_certificate(
        base_pdf_path=base_pdf_path,
        cache_key=cache_key,
        certificate_builder=_builder,
    )
