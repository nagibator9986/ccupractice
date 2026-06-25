"""Generate contract DOCX and PDF files from the configured template."""
from __future__ import annotations

import io
import shutil
import subprocess
import tempfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docxtpl import DocxTemplate
from flask import current_app

from ..models import Contract, CollegeSettings, Partner, Student
from ..utils.files import safe_filename, ensure_dir
from ..utils.time import utc_today
from .qr import generate_qr_png, public_verify_url


_MONTHS_RU = [
    "", "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]

# Brand colours sampled from the CCU logo (coral + charcoal).
_BRAND_CORAL = RGBColor(0xE8, 0x5A, 0x3F)
_BRAND_CHARCOAL = RGBColor(0x46, 0x48, 0x4B)


def _fmt_date(value: date | None) -> str:
    if not value:
        return "____________"
    return value.strftime("%d.%m.%Y")


def _build_context(contract: Contract) -> dict:
    settings: CollegeSettings = CollegeSettings.query.first()
    partner: Partner = contract.partner
    student: Student = contract.student

    d = contract.contract_date or utc_today()
    return {
        "contract": {
            "number": contract.number,
            "date": _fmt_date(d),
            "date_day": f"{d.day:02d}",
            "date_month": _MONTHS_RU[d.month],
            "date_year": d.year,
            "practice_start": _fmt_date(student.practice_start),
            "practice_end": _fmt_date(student.practice_end),
        },
        "college": settings.to_dict() if settings else {},
        "partner": {
            "organization_name": partner.organization_name or "",
            "bin": partner.bin or "",
            "legal_address": partner.legal_address or "",
            "actual_address": partner.actual_address or "",
            "director_full_name": partner.director_full_name or "",
            "director_position": partner.director_position or "Директор",
            "director_basis": partner.director_basis or "Устава",
            "bank_name": partner.bank_name or "",
            "bank_bik": partner.bank_bik or "",
            "bank_iik": partner.bank_iik or "",
            "email": partner.email or "",
            "phone": partner.phone or "",
        },
        "student": {
            "full_name": student.full_name or "",
            "iin": student.iin or "",
            "group_name": student.group_name or "",
            "specialty": student.specialty or "",
            "specialty_code": student.specialty_code or "",
            "course": student.course or "",
            "education_program": student.education_program or student.specialty or "",
            "enrollment_year": student.enrollment_year or (d.year - (student.course or 1) + 1),
            "form_of_study": student.form_of_study or "очная",
            "practice_type": student.practice_type or "профессиональной",
            "practice_start": _fmt_date(student.practice_start),
            "practice_end": _fmt_date(student.practice_end),
            "birth_date": _fmt_date(student.birth_date),
            "id_card_number": student.id_card_number or "",
            "id_card_issued_by": student.id_card_issued_by or "",
            "home_address": student.home_address or "",
            "phone": student.phone or "",
            "legal_rep_full_name": student.legal_rep_full_name or "—",
            "legal_rep_iin": student.legal_rep_iin or "",
            "legal_rep_phone": student.legal_rep_phone or "",
        },
    }


def _template_path() -> Path:
    settings: CollegeSettings = CollegeSettings.query.first()
    template_name = (settings.template_path if settings else "contract_template.docx") or "contract_template.docx"
    base = Path(current_app.config["TEMPLATES_FOLDER"])
    path = base / template_name
    if not path.exists():
        from .template_builder import ensure_default_template
        ensure_default_template(base)
        path = base / "contract_template.docx"
    return path


def _archive_dir(contract: Contract) -> Path:
    base = Path(current_app.config["ARCHIVE_FOLDER"]) / "Профессиональная практика"
    base = base / str(contract.year) / safe_filename(contract.partner.organization_name) / "Договоры"
    return ensure_dir(base)


def _archive_filename(contract: Contract, ext: str) -> str:
    base = f"Договор_практика_{safe_filename(contract.partner.organization_name)}_{safe_filename(contract.student.full_name)}_{contract.year}"
    return f"{base}.{ext}"


# ────────────────────────────────────────────────────────────────
# Verification stamp (QR + footer) — appended after the body render
# ────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply a solid background fill to a DOCX table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_page_number_field(paragraph) -> None:
    """Insert a `PAGE` field at the end of a paragraph (Word page numbering)."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def _footer_band(doc: Document, *, brand_line: str) -> None:
    """Doodocs-style verification band at the bottom of every page.

    Generic over the document type: ``brand_line`` is the right-of-bullet text
    (e.g. ``"CCU PRACTICUM · договор № X"`` or ``"CCU PRACTICUM · LMS-договор № Y"``).
    """
    for section in doc.sections:
        # Drop any prior generated footer band so re-runs do not duplicate it.
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p_elem = p._element
            p_elem.getparent().remove(p_elem)
        for tbl in list(footer.tables):
            tbl._element.getparent().remove(tbl._element)

        tbl = footer.add_table(rows=1, cols=2, width=Cm(17))
        tbl.autofit = True
        left, right = tbl.rows[0].cells
        left.width = Cm(14)
        right.width = Cm(3)

        # Left cell: brand line.
        _set_cell_shading(left, "F1F4F6")
        lp = left.paragraphs[0]
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        run = lp.add_run("● ")
        run.font.color.rgb = _BRAND_CORAL
        run.font.size = Pt(10)
        run.font.bold = True
        run = lp.add_run("Визуализация электронного документа · ")
        run.font.size = Pt(8)
        run.font.color.rgb = _BRAND_CHARCOAL
        run = lp.add_run(brand_line)
        run.font.size = Pt(8)
        run.font.color.rgb = _BRAND_CHARCOAL
        run.font.bold = True

        # Right cell: "стр. N".
        _set_cell_shading(right, "F1F4F6")
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(0)
        rp.paragraph_format.space_after = Pt(0)
        run = rp.add_run("стр. ")
        run.font.size = Pt(8)
        run.font.color.rgb = _BRAND_CHARCOAL
        _add_page_number_field(rp)
        # Style the page-number run too.
        for r in rp.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = _BRAND_CHARCOAL


def _verification_block(doc: Document, *, verify_url: str, verification_code: str) -> None:
    """Append a final "Проверка подлинности" block with a large QR code.

    Layout: section break → header "Список подписей" → 2-col table
    [QR image] [explanation + verify URL + verification code]. Mirrors the
    doodocs.kz signed-document visualization page.
    """
    doc.add_section()

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h = head.add_run("Список подписей")
    h.bold = True
    h.font.size = Pt(20)
    h.font.color.rgb = _BRAND_CHARCOAL

    intro = doc.add_paragraph()
    ir = intro.add_run(
        "Документ согласно п. 1 ст. 7 ЗРК от 7 января 2003 года № 370-II "
        "«Об электронном документе и электронной цифровой подписи» равнозначен "
        "документу на бумажном носителе при наличии действительных ЭЦП всех сторон."
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = _BRAND_CHARCOAL

    # QR + caption table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    qr_cell, info_cell = tbl.rows[0].cells
    qr_cell.width = Cm(4.5)
    info_cell.width = Cm(12)

    # Generate QR PNG (≈ 3.5 cm printed)
    png = generate_qr_png(verify_url, box_size=10, border=2)
    qp = qr_cell.paragraphs[0]
    qp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    qp.add_run().add_picture(io.BytesIO(png), width=Cm(3.8))

    # Info column
    info_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    ip = info_cell.paragraphs[0]
    r = ip.add_run("Проверить подлинность подписания")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = _BRAND_CHARCOAL

    p = info_cell.add_paragraph()
    r = p.add_run(
        "Отсканируйте QR-код камерой смартфона или откройте ссылку:"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = _BRAND_CHARCOAL

    p = info_cell.add_paragraph()
    r = p.add_run(verify_url)
    r.font.size = Pt(10)
    r.font.color.rgb = _BRAND_CORAL
    r.font.bold = True

    p = info_cell.add_paragraph()
    r = p.add_run(f"Код документа: {verification_code}")
    r.font.size = Pt(9)
    r.font.color.rgb = _BRAND_CHARCOAL
    r.italic = True

    p = info_cell.add_paragraph()
    r = p.add_run(
        "Страница проверки покажет всех подписантов, ФИО / ИИН и хэш SHA-256 "
        "подписанного документа в режиме реального времени."
    )
    r.font.size = Pt(9)
    r.font.color.rgb = _BRAND_CHARCOAL


def apply_verification_stamp(
    docx_path: Path | str,
    *,
    brand_line: str,
    verification_code: str,
    verify_url: str,
) -> None:
    """Open the rendered DOCX, embed the per-page footer + final QR block, save in-place.

    Generic — used by both the practicum contract generator and the standalone
    LMS-contract generator. ``brand_line`` is what appears next to the bullet in
    the page footer (e.g. ``"CCU PRACTICUM · договор № X"``); ``verification_code``
    is printed in plain text under the QR for OCR fallback.
    """
    doc = Document(str(docx_path))
    _footer_band(doc, brand_line=brand_line)
    _verification_block(doc, verify_url=verify_url, verification_code=verification_code)
    doc.save(str(docx_path))


def _apply_verification_stamp(docx_path: Path, contract: Contract, verify_url: str) -> None:
    """Practicum-contract backward-compatible wrapper."""
    apply_verification_stamp(
        docx_path,
        brand_line=f"CCU PRACTICUM · договор № {contract.number}",
        verification_code=contract.verification_code,
        verify_url=verify_url,
    )


def _convert_to_pdf(docx_path: Path) -> Path | None:
    """Best-effort DOCX -> PDF conversion using LibreOffice (`soffice`).

    Each call runs with a throwaway, isolated LibreOffice user profile
    (``-env:UserInstallation``). Gunicorn runs several workers/threads that share
    one ``$HOME``; without an isolated profile concurrent conversions serialise
    on (or abort against) the ``~/.config/libreoffice`` lockfile, intermittently
    dropping PDFs under load. Failure modes are logged distinctly instead of
    being silently swallowed so the cause is diagnosable from the server log.
    """
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        current_app.logger.warning(
            "PDF conversion skipped: neither 'soffice' nor 'libreoffice' found on PATH"
        )
        return None
    try:
        with tempfile.TemporaryDirectory(prefix="lo_profile_") as profile_dir:
            subprocess.run(
                [
                    soffice,
                    "-env:UserInstallation=" + Path(profile_dir).as_uri(),
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(docx_path.parent),
                    str(docx_path),
                ],
                check=True,
                timeout=120,
                capture_output=True,
                text=True,
            )
    except subprocess.TimeoutExpired:
        current_app.logger.error("PDF conversion timed out after 120s for %s", docx_path)
        return None
    except subprocess.CalledProcessError as e:
        current_app.logger.error(
            "PDF conversion failed (exit %s) for %s: %s",
            e.returncode, docx_path, (e.stderr or "").strip(),
        )
        return None
    except Exception:
        current_app.logger.exception("PDF conversion crashed for %s", docx_path)
        return None
    return pdf_path if pdf_path.exists() else None


def generate_contract_files(contract: Contract, *, public_base: str | None = None) -> Contract:
    """Render → stamp with QR → convert to PDF.

    `public_base` is the absolute origin the QR-code URL should point at
    (typically the request's `X-Public-Origin` header). If omitted, falls back
    to env `PUBLIC_BASE_URL`; if neither is available, the QR encodes the
    relative path `/verify/<code>`.
    """
    template = _template_path()
    archive_dir = _archive_dir(contract)

    docx_name = _archive_filename(contract, "docx")
    docx_full = archive_dir / docx_name

    tpl = DocxTemplate(str(template))
    tpl.render(_build_context(contract))
    tpl.save(str(docx_full))

    # Embed verification stamp on every page + final QR block — this rewrites
    # the file before any signing happens, so the signed payload always
    # includes the QR + URL (a signer sees what they sign).
    verify_url = public_verify_url(contract.verification_code, base_url=public_base)
    try:
        _apply_verification_stamp(docx_full, contract, verify_url)
    except Exception:
        current_app.logger.exception(
            "Verification stamp embedding failed for contract %s", contract.id
        )

    contract.docx_path = str(docx_full.relative_to(current_app.config["ARCHIVE_FOLDER"]))

    # The DOCX was just rewritten, so any previously produced PDF is now stale.
    contract.pdf_path = None
    stale_pdf = docx_full.with_suffix(".pdf")
    try:
        if stale_pdf.is_file():
            stale_pdf.unlink()
    except OSError as exc:
        current_app.logger.warning("Failed to remove stale PDF %s: %s", stale_pdf, exc)

    produced_pdf = _convert_to_pdf(docx_full)
    if produced_pdf:
        contract.pdf_path = str(produced_pdf.relative_to(current_app.config["ARCHIVE_FOLDER"]))

    return contract
