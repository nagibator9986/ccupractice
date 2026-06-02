"""Programmatically build a docxtpl-ready contract template.

The template mirrors the three-sided practice agreement (Russian only) from the
source "Договор 3-х сторон". Bilingual variant can be added later. Placeholders
follow Jinja2 syntax expected by docxtpl.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


_DEFAULT_FILENAME = "contract_template.docx"


def _add_heading(doc: Document, text: str, *, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def _add_p(doc: Document, text: str, *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def build_template(output_path: str | Path) -> Path:
    output_path = Path(output_path)
    doc = Document()

    # Margins
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(1.8)

    _add_heading(doc, "Договор на проведение профессиональной практики № {{ contract.number }}")
    _add_p(
        doc,
        "город {{ college.city }}                                          «{{ contract.date_day }}» {{ contract.date_month }} {{ contract.date_year }} года",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    _add_p(
        doc,
        "Колледж – Учреждения образования «{{ college.name_ru }}», именуемый в дальнейшем "
        "«организация образования» в лице директора {{ college.director_full_name }}, "
        "действующего на основании {{ college.director_basis }}, с одной стороны,",
    )
    _add_p(
        doc,
        "{{ partner.organization_name }}, БИН {{ partner.bin }}, "
        "юридический адрес: {{ partner.legal_address }}, именуемый в дальнейшем «предприятие», "
        "в лице {{ partner.director_position }} {{ partner.director_full_name }}, "
        "действующего на основании {{ partner.director_basis }}, с другой стороны,",
    )
    _add_p(
        doc,
        "и гражданин(ка) {{ student.full_name }}, ИИН {{ student.iin }}, "
        "именуемый(ая) в дальнейшем «обучающийся», с третьей стороны, "
        "в соответствии с действующим законодательством Республики Казахстан, заключили настоящий договор о нижеследующем:",
    )

    _add_heading(doc, "1. ПРЕДМЕТ ДОГОВОРА", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_p(
        doc,
        "1. Организация образования осуществляет обучение обучающегося, "
        "поступившего в {{ student.enrollment_year }} году, по образовательной программе "
        "«{{ student.education_program }}», специальности (квалификации) "
        "{{ student.specialty_code }} «{{ student.specialty }}».",
    )
    _add_p(
        doc,
        "2. Предприятие обеспечивает обучающегося базой профессиональной практики в соответствии "
        "с профилем образовательной программы на безвозмездной основе.",
    )
    _add_p(
        doc,
        "3. Обучающийся осваивает образовательную программу с целью получения ключевых и "
        "профессиональных компетенций, позволяющих квалифицированно выполнять производственные функции и задачи.",
    )

    _add_heading(doc, "2. ПРАВА И ОБЯЗАННОСТИ СТОРОН", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_p(doc, "4. Организация образования обязуется:", bold=True)
    obligations_college = [
        "направить на предприятие обучающегося по специальности {{ student.specialty_code }} «{{ student.specialty }}», "
        "формы обучения {{ student.form_of_study }} для прохождения {{ student.practice_type }} практики "
        "в соответствии с академическим календарем в период с {{ contract.practice_start }} по {{ contract.practice_end }};",
        "ознакомить обучающегося с его обязанностями и ответственностью, указанных в настоящем Договоре;",
        "разработать и согласовать с предприятием программу профессиональной практики и календарные графики прохождения;",
        "за две недели до начала практики предоставить предприятию программу и календарные графики с указанием количества обучающихся;",
        "назначить приказом руководителей практики из числа преподавателей соответствующих специальностей;",
        "обеспечить соблюдение обучающимся трудовой дисциплины и правил внутреннего распорядка предприятия;",
        "организовать прохождение и осуществлять периодический контроль профессиональной практики в соответствии с образовательной программой;",
        "оказывать работникам предприятия методическую помощь в организации профессиональной практики;",
        "при необходимости предоставлять предприятию сведения об учебных достижениях обучающегося;",
        "принимать участие в расследовании несчастных случаев с участием обучающегося;",
        "в случае ликвидации поставить в известность предприятие и принять меры к переводу обучающегося;",
        "при необходимости предусмотреть жилищно-бытовые и другие условия для обучающегося.",
    ]
    for i, item in enumerate(obligations_college, start=1):
        _add_p(doc, f"{i}) {item}")

    _add_p(doc, "5. Организация образования имеет право:", bold=True)
    _add_p(
        doc,
        "1) расторгнуть Договор в одностороннем порядке при самовольном прекращении обучения, "
        "оставлении на повторный год обучения, а также при отчислении обучающегося в порядке, "
        "определенном законодательством Республики Казахстан.",
    )

    _add_p(doc, "6. Предприятие обязуется:", bold=True)
    obligations_partner = [
        "обеспечить обучающемуся условия безопасной работы на рабочем месте, провести инструктажи и при необходимости обучение безопасным методам труда;",
        "принять по направлению на профессиональную практику обучающегося в соответствии с условиями настоящего договора;",
        "не допускать использования обучающегося на должностях, не предусмотренных программой практики;",
        "обеспечить квалифицированных специалистов для руководства профессиональной практикой;",
        "сообщать в организацию образования о всех случаях нарушения обучающимся дисциплины и правил внутреннего распорядка;",
        "создать необходимые условия для выполнения программы практики и индивидуальных заданий;",
        "по окончании практики выдать характеристику о работе обучающегося и выставить оценку.",
    ]
    for i, item in enumerate(obligations_partner, start=1):
        _add_p(doc, f"{i}) {item}")

    _add_p(doc, "7. Предприятие имеет право:", bold=True)
    rights_partner = [
        "участвовать в разработке образовательной программы профессиональной практики;",
        "предлагать темы курсовых и дипломных работ в соответствии с потребностями предприятия;",
        "принимать участие в итоговой аттестации обучающихся;",
        "запрашивать информацию о текущей успеваемости обучающихся;",
        "требовать от организации образования качественного обучения обучающихся;",
        "рассматривать кандидатуру выпускника, обучавшегося по образовательному гранту, для приёма на работу.",
    ]
    for i, item in enumerate(rights_partner, start=1):
        _add_p(doc, f"{i}) {item}")

    _add_p(doc, "8. Обучающийся обязан:", bold=True)
    obligations_student = [
        "соблюдать трудовую дисциплину, правила внутреннего распорядка, пожарной безопасности и техники безопасности;",
        "бережно относиться к оборудованию, приборам, документации и другому имуществу предприятия;",
        "строго соблюдать и выполнять требования программы практики;",
        "прибыть в распоряжение предприятия к установленному сроку;",
        "не разглашать конфиденциальную информацию о предприятии в процессе и после прохождения практики.",
    ]
    for i, item in enumerate(obligations_student, start=1):
        _add_p(doc, f"{i}) {item}")

    _add_p(doc, "9. Обучающийся имеет право:", bold=True)
    rights_student = [
        "пользоваться необходимыми инструментами, оборудованием, приборами и материалами по согласованию с наставником;",
        "на возмещение вреда, причинённого здоровью в процессе прохождения профессиональной подготовки;",
        "после успешного прохождения итоговой аттестации продолжить работу по полученной квалификации на предприятии при наличии вакансии.",
    ]
    for i, item in enumerate(rights_student, start=1):
        _add_p(doc, f"{i}) {item}")

    _add_heading(doc, "3. ОТВЕТСТВЕННОСТЬ СТОРОН", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_p(
        doc,
        "10. За неисполнение либо ненадлежащее исполнение своих обязанностей, предусмотренных настоящим "
        "Договором, стороны несут ответственность, установленную действующим законодательством Республики Казахстан.",
    )

    _add_heading(doc, "4. ПОРЯДОК РАЗРЕШЕНИЯ СПОРОВ", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_p(
        doc,
        "11. Разногласия и споры, возникающие в процессе выполнения настоящего Договора, "
        "разрешаются непосредственно сторонами в целях выработки взаимоприемлемых решений.",
    )
    _add_p(
        doc,
        "12. Вопросы, не разрешённые сторонами путём переговоров, разрешаются в соответствии "
        "с действующим законодательством Республики Казахстан.",
    )

    _add_heading(doc, "5. СРОК ДЕЙСТВИЯ И ПОРЯДОК ИЗМЕНЕНИЯ ДОГОВОРА", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_p(doc, "13. Настоящий договор вступает в силу со дня его подписания сторонами и действует до полного его исполнения.")
    _add_p(doc, "14. Условия настоящего Договора могут быть изменены и дополнены по взаимному письменному соглашению сторон.")
    _add_p(doc, "15. Настоящий Договор заключается в трёх экземплярах на государственном и русском языках, имеющих одинаковую юридическую силу.")
    _add_p(doc, "16. Юридические адреса и банковские реквизиты Сторон:", bold=True)

    # Requisites
    _add_p(doc, "Организация образования:", bold=True)
    _add_p(doc, "{{ college.name_ru }}")
    _add_p(doc, "{{ college.address }}")
    _add_p(doc, "БИН – {{ college.bin }}")
    _add_p(doc, "ИИК {{ college.iik }}")
    _add_p(doc, "{{ college.bank_name }}, {{ college.bank_address }}")
    _add_p(doc, "БИК – {{ college.bank_bik }}")
    _add_p(doc, "E-mail: {{ college.email }}")
    _add_p(doc, "Тел: {{ college.phone }}")
    _add_p(doc, "Директор: ____________ {{ college.director_full_name }}")

    _add_p(doc, " ")
    _add_p(doc, "Предприятие:", bold=True)
    _add_p(doc, "{{ partner.organization_name }}")
    _add_p(doc, "{{ partner.legal_address }}")
    _add_p(doc, "БИН {{ partner.bin }}")
    _add_p(doc, "ИИК {{ partner.bank_iik }}")
    _add_p(doc, "{{ partner.bank_name }}")
    _add_p(doc, "БИК {{ partner.bank_bik }}")
    _add_p(doc, "{{ partner.director_position }}: ____________ {{ partner.director_full_name }}")

    _add_p(doc, " ")
    _add_p(doc, "Обучающийся:", bold=True)
    _add_p(doc, "{{ student.full_name }}")
    _add_p(doc, "Дата рождения: {{ student.birth_date }}, ИИН: {{ student.iin }}")
    _add_p(doc, "Удостоверение личности № {{ student.id_card_number }}, выдано {{ student.id_card_issued_by }}")
    _add_p(doc, "Адрес: {{ student.home_address }}, тел.: {{ student.phone }}")
    _add_p(doc, "Подпись: ____________")

    _add_p(doc, " ")
    _add_p(doc, "Законный представитель (при необходимости):", bold=True)
    _add_p(doc, "{{ student.legal_rep_full_name }}, ИИН {{ student.legal_rep_iin }}, тел.: {{ student.legal_rep_phone }}")
    _add_p(doc, "Подпись: ____________")

    doc.save(output_path)
    return output_path


def ensure_default_template(templates_dir: str | Path) -> Path:
    templates_dir = Path(templates_dir)
    templates_dir.mkdir(parents=True, exist_ok=True)
    path = templates_dir / _DEFAULT_FILENAME
    if not path.exists():
        build_template(path)
    return path
