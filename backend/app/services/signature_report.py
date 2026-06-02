"""Build a human-readable PDF/DOCX 'Signature audit report' for a contract.

The report lists each signing party, the cryptographic identity extracted
from the CMS, the SHA-256 hash of the signed payload, and a timestamp.
Combined with the original DOCX it constitutes the legal signing bundle.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from flask import current_app

from ..models import Contract, SigningRequest
from ..utils.files import safe_filename
from ..utils.time import utc_now


CORAL = RGBColor(0xE8, 0x5A, 0x3F)
CHARCOAL = RGBColor(0x46, 0x48, 0x4B)


def _heading(doc: Document, text: str, *, size=18, color=CORAL, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Inter"
    return p


def _line(doc: Document, label: str, value: str, *, bold=False):
    p = doc.add_paragraph()
    a = p.add_run(f"{label}: ")
    a.bold = True
    a.font.color.rgb = CHARCOAL
    a.font.name = "Inter"
    a.font.size = Pt(10)
    b = p.add_run(value or "—")
    b.font.name = "Inter"
    b.font.size = Pt(10)
    if bold:
        b.bold = True


def build_signature_report(contract: Contract) -> Path:
    settings_folder = Path(current_app.config["ARCHIVE_FOLDER"])
    out_dir = settings_folder / "Профессиональная практика" / str(contract.year) / safe_filename(contract.partner.organization_name) / "Подписи"
    out_dir.mkdir(parents=True, exist_ok=True)
    fname = f"Отчёт_подписания_{safe_filename(contract.number)}.docx"
    out_path = out_dir / fname

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    _heading(doc, f"Отчёт о подписании договора № {contract.number}", size=18)
    _heading(doc, "CCU PRACTICUM • College of Caspian University", size=10, color=CHARCOAL)
    doc.add_paragraph()

    _line(doc, "Партнёр", contract.partner.organization_name, bold=True)
    _line(doc, "Студент", contract.student.full_name, bold=True)
    _line(doc, "Дата договора", contract.contract_date.strftime("%d.%m.%Y") if contract.contract_date else "—")
    _line(doc, "Период практики", _format_practice_period(contract))
    _line(doc, "Текущий статус", contract.status)
    _line(doc, "Сформирован отчёт", utc_now().strftime("%d.%m.%Y %H:%M UTC"))

    doc.add_paragraph()
    _heading(doc, "Подписи сторон", size=14, color=CORAL)

    role_labels = {
        "college": "Колледж (организация образования)",
        "partner": "Предприятие (партнёр)",
        "student": "Обучающийся",
    }

    for role in ("college", "partner", "student"):
        _heading(doc, role_labels[role], size=12, color=CHARCOAL)
        signature = next((s for s in contract.signatures if s.signer_role == role), None)
        sr = next((r for r in contract.signing_requests if r.signer_role == role), None)

        if signature:
            _line(doc, "ФИО подписанта", signature.signer_full_name)
            _line(doc, "ИИН/БИН", signature.signer_iin_or_bin)
            _line(doc, "Серийный № сертификата", signature.signer_serial)
            _line(doc, "SHA-256 подписанного документа", signature.signed_payload_sha256)
            _line(doc, "Дата и время подписания (UTC)", signature.created_at.strftime("%d.%m.%Y %H:%M:%S"))
            _line(doc, "Способ подписания", "ЭЦП (CAdES/CMS, NCALayer)")
        else:
            status = sr.status if sr else "не отправлено"
            _line(doc, "Статус", _status_label(status))
            if sr:
                _line(doc, "Получатель", sr.recipient_name or "—")
                _line(doc, "Email", sr.recipient_email or "—")
                _line(doc, "Срок действия ссылки", sr.expires_at.strftime("%d.%m.%Y") if sr.expires_at else "—")
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Документ сформирован автоматически. Каждая подпись является CAdES/CMS-структурой, "
                  "содержащей оригинальный SHA-256 хэш Word-документа и сертификат подписанта, "
                  "выданный НУЦ Республики Казахстан.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = CHARCOAL

    doc.save(out_path)
    return out_path


def _status_label(status: str) -> str:
    return {
        "pending": "Отправлено, ожидает подписи",
        "viewed": "Просмотрено получателем",
        "signed": "Подписано ЭЦП",
        "revoked": "Отозвано администратором",
        "не отправлено": "Ссылка ещё не отправлена",
    }.get(status, status)


def _format_practice_period(contract: Contract) -> str:
    s = contract.student
    if not (s and s.practice_start):
        return "—"
    return f"{s.practice_start.strftime('%d.%m.%Y')} → {s.practice_end.strftime('%d.%m.%Y') if s.practice_end else '?'}"
