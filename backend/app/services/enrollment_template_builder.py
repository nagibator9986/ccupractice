"""Build docxtpl-ready templates for the enrollment documents.

The College's official contracts are **bilingual two-column tables** (Kazakh in
the left column, Russian in the right) with a fixed legal layout and section
order. Earlier this module hand-rebuilt a Russian-only, single-column draft —
which is exactly why the platform rendered the documents in a "chaotic order"
that didn't match the originals. We now instead take the **real source .docx**
(shipped under ``templates_docx/source/``) and inject docxtpl/Jinja tags *only*
into the blank fill-in fields, leaving every byte of the official bilingual
layout, wording and signature placement untouched.

Produced templates (rendered later by :mod:`enrollment_documents`):
  * contract_enrollment_template_v3.docx — «Договор на оказание образовательных
    услуг» (bilingual KZ/RU)
  * contract_lms_template_v1.docx        — «Договор о подключении к цифровой
    экосистеме Caspian College (Caspian Digital)» (bilingual KZ/RU)
  * consent_template_v2.docx             — «Согласие на сбор и обработку
    персональных данных» (Russian, age-aware — matches the RU-only original)

The injected variables (``{{ student.full_name }}`` …) are filled from the
EnrollmentContract / CollegeSettings in :func:`enrollment_documents._build_context`.
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Bump the filename to force regeneration when the injection logic changes
# (ensure_enrollment_templates only builds a file that doesn't already exist).
CONTRACT_FILENAME = "contract_enrollment_template_v5.docx"
LMS_FILENAME = "contract_lms_template_v1.docx"
# v4: rebuilt from scratch to match the official College of Caspian University
# consent form wording word-for-word (was previously producing a shortened
# variant that lacked "рекламных целях / потенциальным работодателям /
# составления соответствующего письменного документа" phrases). The consent
# is now ALSO rebuilt on every ensure_enrollment_templates call — see the
# note in that function — because a stale file on the deployed volume kept
# masking earlier fixes even after the filename was bumped.
CONSENT_FILENAME = "consent_template_v4.docx"

SOURCE_DIRNAME = "source"
SOURCE_CONTRACT = "source_contract_edu.docx"
SOURCE_LMS = "source_contract_lms.docx"

# A "blank" fill-in field: a run of 3+ underscores. Jinja tags carry no
# underscores, so a filled field is never re-matched.
_BLANK = re.compile(r"_{3,}")


def _fill_paragraph(paragraph, tags_by_occ: dict[int, str]) -> int:
    """Replace selected blank fields inside one paragraph with Jinja tags.

    ``tags_by_occ`` maps the 1-based blank index (counted left-to-right across
    ALL runs of the paragraph) to the replacement text. Blanks whose index is
    absent are left as-is (e.g. a signature line we want to keep blank).

    Replacement happens *inside the run that holds the blank*, so the run-level
    character formatting of the rest of the paragraph (bold quoted terms, etc.)
    is fully preserved — and the inserted tag lands in a single run, so docxtpl
    never sees a Jinja tag split across runs.
    """
    occ = 0
    replaced = 0
    for run in paragraph.runs:
        text = run.text
        if not _BLANK.search(text):
            continue
        out, last = [], 0
        for m in _BLANK.finditer(text):
            occ += 1
            out.append(text[last:m.start()])
            if occ in tags_by_occ:
                out.append(tags_by_occ[occ])
                replaced += 1
            else:
                out.append(text[m.start():m.end()])
            last = m.end()
        out.append(text[last:])
        run.text = "".join(out)
    return replaced


def _apply_fills(doc: Document, fills: list[tuple]) -> int:
    """Apply a fill map to the first table of ``doc``.

    Each fill is ``(row, col, paragraph_index, {occ: tag})``. Returns the number
    of fields actually replaced so the builder can assert nothing silently
    drifted out of place after an edit to the source document.
    """
    table = doc.tables[0]
    total = 0
    for row, col, para_idx, tags in fills:
        cell = table.rows[row].cells[col]
        paragraph = cell.paragraphs[para_idx]
        total += _fill_paragraph(paragraph, tags)
    return total


# ─────────────────────────────────────────────────────────────────────────────
# Fill maps — (row, col, paragraph_index, {blank_occurrence: jinja_tag})
# Coordinates are verified against the source .docx; _build_from_source() asserts
# the expected number of replacements so a source change can't silently misalign.
# ─────────────────────────────────────────────────────────────────────────────

_NUM = "{{ contract.number }}"
_NAME = "{{ student.full_name }}"
_IIN = "{{ student.iin }}"
_SPEC = "{{ student.specialty }}"
_QUAL = "{{ student.qualification }}"
_AMOUNT = " {{ contract.tuition_amount }} "
_CITY = "{{ student.addr_city }}"
_DISTRICT = "{{ student.addr_district }}"
_STREET = " {{ student.addr_street }}"
_HOUSE = "{{ student.addr_house }}"
_DOCNO = "{{ student.id_doc_number }}"
_DOCBY = "{{ student.id_doc_issued_by }}"
_DOCDATE = "{{ student.id_doc_issued_date }}"
_HOMEPHONE = "{{ student.home_phone }}"
_PHONE = "{{ student.phone }}"
_PNAME = "{{ parent.full_name }}"
_PADDR = "{{ parent.address }}"
_PPHONE = "{{ parent.phone }}"
_PEMAIL = "{{ parent.email }}"


# Common bottom block: «СТУДЕНТ» requisites + «СОГЛАСИЕ РОДИТЕЛЕЙ». Row offsets
# differ between the two contracts, so each contract passes its own absolute map.
EDU_FILLS = [
    (0, 0, 0, {1: _NUM}), (0, 1, 0, {1: _NUM}),                 # № договора (KZ/RU header)
    (2, 0, 0, {1: _NAME}), (2, 1, 0, {1: _NAME}),               # ФИО студента (преамбула)
    (4, 0, 0, {1: _SPEC}), (4, 1, 0, {1: _SPEC}),               # специальность (предмет)
    (12, 0, 0, {1: _NAME}), (12, 1, 0, {1: _NAME}),             # ФИО (обязанности колледжа)
    (12, 0, 7, {1: _QUAL}), (12, 1, 7, {1: _QUAL}),             # квалификация
    (16, 0, 0, {1: _AMOUNT}), (16, 1, 0, {1: _AMOUNT}),         # стоимость за год
    (27, 1, 1, {1: _NAME}),                                     # ФИО студента (реквизиты)
    (28, 1, 0, {1: _IIN}),                                      # ИИН
    (29, 1, 1, {1: _CITY}), (29, 1, 2, {1: _DISTRICT}),         # адрес: город / район
    (29, 1, 3, {1: _STREET, 2: _HOUSE}),                        # улица / дом
    (30, 1, 1, {1: _DOCNO}), (30, 1, 2, {1: _DOCBY}),           # удостоверение: №, кем выдано
    (30, 1, 3, {1: _DOCDATE}),                                  # дата выдачи
    (30, 1, 4, {1: _HOMEPHONE}), (30, 1, 5, {1: _PHONE}),       # тел. дом / сот
    (32, 1, 2, {2: _NAME}),                                     # подпись: оставить линию, ФИО справа
    (34, 0, 0, {1: _PNAME}), (34, 0, 3, {1: _PADDR}),           # родитель: ФИО, адрес (KZ)
    (34, 0, 5, {1: _PPHONE}), (34, 0, 7, {1: _PEMAIL}),         # родитель: тел, email (KZ)
    (34, 1, 0, {1: _PNAME}), (34, 1, 2, {1: _PADDR}),           # родитель: ФИО, адрес (RU)
    (34, 1, 4, {1: _PPHONE}), (34, 1, 6, {1: _PEMAIL}),         # родитель: тел, email (RU)
]

LMS_FILLS = [
    (0, 0, 0, {1: _NUM}), (0, 1, 0, {1: _NUM}),                 # № договора (KZ/RU header)
    (2, 0, 0, {1: _NAME}), (2, 1, 0, {1: _NAME}),               # ФИО студента (преамбула)
    (26, 1, 1, {1: _NAME}),                                     # ФИО студента (реквизиты)
    (27, 1, 0, {1: _IIN}),                                      # ИИН
    (28, 1, 1, {1: _CITY}), (28, 1, 2, {1: _DISTRICT}),         # адрес: город / район
    (28, 1, 3, {1: _STREET, 2: _HOUSE}),                        # улица / дом
    (29, 1, 1, {1: _DOCNO}), (29, 1, 2, {1: _DOCBY}),           # удостоверение: №, кем выдано
    (29, 1, 3, {1: _DOCDATE}),                                  # дата выдачи
    (29, 1, 4, {1: _HOMEPHONE}), (29, 1, 5, {1: _PHONE}),       # тел. дом / сот
    (31, 1, 2, {2: _NAME}),                                     # подпись: оставить линию, ФИО справа
    (33, 0, 0, {1: _PNAME}), (33, 0, 3, {1: _PADDR}),           # родитель: ФИО, адрес (KZ)
    (33, 0, 5, {1: _PPHONE}), (33, 0, 7, {1: _PEMAIL}),         # родитель: тел, email (KZ)
    (33, 1, 0, {1: _PNAME}), (33, 1, 2, {1: _PADDR}),           # родитель: ФИО, адрес (RU)
    (33, 1, 4, {1: _PPHONE}), (33, 1, 6, {1: _PEMAIL}),         # родитель: тел, email (RU)
]

# Expected replacement counts — a guard against a source-file edit silently
# shifting paragraph indices (the build then raises instead of producing a
# subtly-broken template).
_EDU_EXPECTED = sum(len(t) for *_, t in EDU_FILLS)
_LMS_EXPECTED = sum(len(t) for *_, t in LMS_FILLS)


def _build_from_source(source_path: Path, fills: list[tuple], expected: int,
                       out_path: Path) -> Path:
    doc = Document(str(source_path))
    replaced = _apply_fills(doc, fills)
    if replaced != expected:
        raise RuntimeError(
            f"Template injection mismatch for {source_path.name}: "
            f"replaced {replaced} fields, expected {expected}. "
            "The source .docx layout changed — re-verify the fill map."
        )
    # Atomic write: save to a temp path in the same directory, then rename onto
    # the target. Guarantees a concurrent Flask worker never sees a half-written
    # .docx (python-docx does a streaming ZIP write, so an interrupted save
    # leaves a corrupt archive — docxtpl then raises on next render). The
    # rename() is atomic within a filesystem, so worker A reads the OLD file
    # in full, worker B reads the NEW file in full — never a partial one.
    out_path = Path(out_path)
    tmp_path = out_path.with_suffix(out_path.suffix + ".tmp")
    try:
        doc.save(str(tmp_path))
        tmp_path.replace(out_path)
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass
    return out_path


def build_contract_template(out_path: str | Path, source_dir: str | Path) -> Path:
    source = Path(source_dir) / SOURCE_CONTRACT
    return _build_from_source(source, EDU_FILLS, _EDU_EXPECTED, Path(out_path))


def build_lms_template(out_path: str | Path, source_dir: str | Path) -> Path:
    source = Path(source_dir) / SOURCE_LMS
    return _build_from_source(source, LMS_FILLS, _LMS_EXPECTED, Path(out_path))


# ─────────────────────────────────────────────────────────────────────────────
# Personal-data consent (Russian, age-aware) — the original consent is a
# single-column Russian document, so this one stays single-language by design.
# ─────────────────────────────────────────────────────────────────────────────

def _doc() -> Document:
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.6)
    s.bottom_margin = Cm(1.6)
    s.left_margin = Cm(2.0)
    s.right_margin = Cm(1.6)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    return doc


def _ctrl(doc, tag: str):
    doc.add_paragraph(tag)


def _h(doc, text, *, size=12, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    return p


def _p(doc, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    return p


def build_consent_template(path: str | Path) -> Path:
    """Programmatically build the personal-data consent template.

    Wording is the OFFICIAL College of Caspian University consent form —
    matches the printed original word-for-word so the legal fidelity of the
    document is preserved. Underscore fill-in placeholders from the printed
    form are replaced with docxtpl ``{{ jinja }}`` tags that get filled from
    the enrollment record at generation time.
    """
    path = Path(path)
    doc = _doc()

    _h(doc, "СОГЛАСИЕ")
    _h(doc, "на сбор и обработку персональных данных", size=12)

    # ── Preamble — age-aware. Uses the official form's wording exactly. ─────
    _ctrl(doc, "{%p if applicant.is_minor %}")
    _p(
        doc,
        "Я, несовершеннолетний(-яя) {{ applicant.full_name }} (ФИО), "
        "дата рождения {{ applicant.birth_date }} удостоверение личности/свидетельство о рождении "
        "номер {{ applicant.id_doc_number }}, выданное {{ applicant.id_doc_issued_by }} "
        "(кем и когда), зарегистрированный(-ая) по адресу: {{ applicant.address }} "
        "(далее – «Обучающийся»), действующий(-ая) с согласия законного представителя "
        "{{ parent.full_name }} (ФИО) удостоверение личности: номер {{ parent.id_doc_number }}, "
        "выданное {{ parent.id_doc_issued_by }} (кем и когда), зарегистрированный по адресу: "
        "{{ parent.address }},",
    )
    _ctrl(doc, "{%p else %}")
    _p(
        doc,
        "Я, {{ applicant.full_name }} (ФИО), дата рождения {{ applicant.birth_date }} "
        "удостоверение личности номер {{ applicant.id_doc_number }}, выданное "
        "{{ applicant.id_doc_issued_by }} (кем и когда), зарегистрированный(-ая) по адресу: "
        "{{ applicant.address }} (далее – «Обучающийся»),",
    )
    _ctrl(doc, "{%p endif %}")

    # ── Operator + processing scope — matches the official form's wording ──
    _p(
        doc,
        "даю согласие оператору – Учреждению Образования «{{ college.name_ru }}» "
        "БИН {{ college.bin }}, адрес: {{ college.address }} на сбор и обработку, "
        "включая, но не ограничиваясь: на систематизацию, накопление, хранение, "
        "уточнение (обновление, изменение, дополнение), использование, обезличивание, "
        "блокирование, уничтожение, передачу другим лицам следующих персональных данных:",
        bold=True,
    )
    for item in [
        "Фамилия/Имя/Отчество.",
        "ИИН.",
        "Дата рождения.",
        "Пол.",
        "Номер мобильного телефона.",
        "Адрес электронной почты (email).",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"

    # ── Purpose — verbatim from the official form ──────────────────────────
    _p(
        doc,
        "\tЦель сбора и обработки: предоставление Обучающемуся и/или его "
        "Законным представителям, сотрудникам организаций системы образования "
        "информации о текущей успеваемости Обучающегося в образовательной "
        "организации в электронном формате, обеспечение процессов оказания "
        "государственных услуг в электронном виде в сфере образования, сбор "
        "обезличенных данных по успеваемости для статистических исследований, "
        "использование данных в рекламных целях, передачу данных потенциальным "
        "работодателям, а также для ведения статистики и учета.",
    )
    _p(
        doc,
        "\tНастоящее согласие в отношении сбора и обработки указанных данных "
        "действует на весь период обучения Обучающегося в указанной образовательной "
        "организации до момента выпуска, исключения, перевода в другую "
        "образовательную организацию.",
    )
    _p(
        doc,
        "\tДаю свое согласие на хранение указанных персональных данных в "
        "соответствующих архивах Оператора в течение срока, установленного "
        "законодательством Республики Казахстан.",
    )
    _p(
        doc,
        "\tОсведомлен(а) о праве отозвать свое согласие посредством составления "
        "соответствующего письменного документа, который может быть направлен мной "
        "на адрес образовательной организации по почте заказным письмом с уведомлением "
        "о вручении, либо вручен лично под расписку представителю образовательной "
        "организации.",
    )

    _p(doc, " ")
    _ctrl(doc, "{%p if applicant.is_minor %}")
    _p(doc, "Подпись несовершеннолетнего: ____________________ / {{ applicant.full_name }}")
    _ctrl(doc, "{%p else %}")
    _p(doc, "Подпись Обучающегося: ____________________ / {{ applicant.full_name }}")
    _ctrl(doc, "{%p endif %}")
    _p(doc, "Согласна (согласен) ____________________ (подпись законного представителя) / {{ parent.full_name }}")
    _p(doc, "Дата ____________")

    # Atomic write — see the note in _build_from_source for the rationale.
    tmp_path = path.with_suffix(path.suffix + ".tmp")
    try:
        doc.save(str(tmp_path))
        tmp_path.replace(path)
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass
    return path


# ─────────────────────────────────────────────────────────────────────────────

def ensure_enrollment_templates(templates_dir: str | Path) -> dict:
    """Build any missing template into ``templates_dir`` and return their paths.

    Sources live in ``templates_dir/source/``. The bilingual contracts are
    injected from those sources; the consent is generated programmatically.

    IMPORTANT: the consent template is ALWAYS rebuilt on every call because
    it is fully programmatic (no source .docx dependency, cheap to build)
    and we've been bitten by stale files persisting on Railway's volume
    across deploys — a filename bump doesn't help if any variant of the
    file is already sitting on disk. Rebuilding on every generate call
    guarantees the deployed code is the single source of truth for consent
    layout + Jinja tags. Contracts stay build-once because they depend on
    a source .docx that ships with the code.
    """
    templates_dir = Path(templates_dir)
    templates_dir.mkdir(parents=True, exist_ok=True)
    source_dir = templates_dir / SOURCE_DIRNAME

    contract = templates_dir / CONTRACT_FILENAME
    lms = templates_dir / LMS_FILENAME
    consent = templates_dir / CONSENT_FILENAME

    # Build CONSENT first — it has no source-file dependency and is cheap to
    # rebuild. Sweep any older-version consent files so a stale template that
    # persisted through past deploys can't shadow the current one.
    for stale in templates_dir.glob("consent_template_v*.docx"):
        if stale.resolve() != consent.resolve():
            try:
                stale.unlink()
            except OSError:
                pass
    build_consent_template(consent)

    # Contract + LMS: always rebuild too. Both depend on source .docx files
    # that ship with the code, but the built template file itself is a
    # derivative that has already been bitten by the "stale on Railway
    # volume" failure mode (see the consent history above). Cheap to rebuild
    # (~milliseconds via python-docx), so no reason to skip.
    for stale in templates_dir.glob("contract_enrollment_template_v*.docx"):
        if stale.resolve() != contract.resolve():
            try:
                stale.unlink()
            except OSError:
                pass
    if (source_dir / SOURCE_CONTRACT).is_file():
        build_contract_template(contract, source_dir)

    for stale in templates_dir.glob("contract_lms_template_v*.docx"):
        if stale.resolve() != lms.resolve():
            try:
                stale.unlink()
            except OSError:
                pass
    if (source_dir / SOURCE_LMS).is_file():
        build_lms_template(lms, source_dir)

    return {"contract": contract, "lms": lms, "consent": consent}
