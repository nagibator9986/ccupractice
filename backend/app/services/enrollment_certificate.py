"""Generate the standalone "Сертификат подписания" PDF for an enrollment.

Why a separate document?
========================
The originally-signed DOCX bytes carry the legal signatures (the CMS payload
is bound to ``SHA-256(docx_bytes)``). We CANNOT add a signature visualization
to the signed bytes after the fact — that would invalidate every existing
``EnrollmentSignature`` row.

Instead we generate a SEPARATE, unsigned PDF on demand. It lists every signer,
their identity, certificate serial, signing time and verification level, and
embeds the same QR code that links to the public ``/verify/<code>`` page. This
mirrors how doodocs.kz / eGov.kz visualize signed documents: the legal payload
stays untouched, the visualization is an annex that anyone can regenerate.

SAFETY GUARANTEES
=================
- This function NEVER touches ``contract_docx_path`` / ``consent_docx_path``
  / ``contract_pdf_path`` / ``consent_pdf_path`` on the enrollment.
- It NEVER touches any ``EnrollmentSignature`` row.
- The output is written to a separate ``Сертификаты_подписания/`` subtree —
  it lives next to the legal documents but is logically a side-channel artefact.
- It can be regenerated any number of times without consequence.
"""
from __future__ import annotations

import io
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from flask import current_app

from ..models import (
    CollegeSettings,
    DOCUMENT_LABELS,
    EnrollmentContract,
    PARTY_COLLEGE,
    PARTY_LABELS,
)
from ..utils.files import ensure_dir, safe_filename
from ..utils.time import utc_now
from .document_generator import _convert_to_pdf
from .qr import generate_qr_png, public_verify_url


_BRAND_CORAL = RGBColor(0xE8, 0x5A, 0x3F)
_BRAND_CHARCOAL = RGBColor(0x46, 0x48, 0x4B)
_BRAND_INK = RGBColor(0x1F, 0x20, 0x24)

_LEVEL_RU = {
    "legal":          "Юридически верифицирована (NCANode)",
    "full":           "Криптографически проверена",
    "document_bound": "Привязана к документу (GOST)",
    "accepted":       "Принята",
}


def _mask_iin(value: str | None) -> str:
    if not value:
        return "—"
    digits = "".join(ch for ch in str(value) if ch.isdigit())
    if len(digits) < 8:
        return digits or "—"
    return f"{digits[:6]}****{digits[-2:]}"


def _serial_tail(value: str | None) -> str:
    if not value:
        return "—"
    return "…" + str(value)[-12:].upper()


def _archive_dir(e: EnrollmentContract) -> Path:
    base = Path(current_app.config["ARCHIVE_FOLDER"]) / "Сертификаты_подписания"
    base = base / str(e.year or (e.contract_date or date.today()).year)
    base = base / safe_filename(e.applicant_full_name or f"enroll_{e.id}")
    return ensure_dir(base)


def _output_paths(e: EnrollmentContract) -> tuple[Path, Path]:
    """Returns ``(docx_path, pdf_path)`` for the certificate output."""
    who = safe_filename(e.applicant_full_name or f"enroll_{e.id}")
    year = e.year or (e.contract_date or date.today()).year
    archive_dir = _archive_dir(e)
    stem = f"Сертификат_подписания_{who}_{year}"
    return archive_dir / f"{stem}.docx", archive_dir / f"{stem}.pdf"


# ────────────────────────────────────────────────────────────────────────────
# DOCX writers — small helpers
# ────────────────────────────────────────────────────────────────────────────

def _h(doc, text, *, size=18, color=_BRAND_CORAL, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Inter"
    return p


def _p(doc, text, *, size=10, color=_BRAND_INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
       italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Inter"
    r.bold = bold
    r.italic = italic
    return p


def _kv(doc, label, value, *, bold_value=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    rl = p.add_run(f"{label}: ")
    rl.bold = True
    rl.font.color.rgb = _BRAND_CHARCOAL
    rl.font.size = Pt(10)
    rl.font.name = "Inter"
    rv = p.add_run(value if value not in (None, "") else "—")
    rv.font.size = Pt(10)
    rv.font.color.rgb = _BRAND_INK
    rv.bold = bold_value
    rv.font.name = "Inter"


# ────────────────────────────────────────────────────────────────────────────
# Public API
# ────────────────────────────────────────────────────────────────────────────

def generate_signature_certificate(
    e: EnrollmentContract,
    *,
    public_base: str | None = None,
) -> tuple[Path, Path | None]:
    """Render the "Сертификат подписания" DOCX + PDF for `e`.

    Returns ``(docx_path, pdf_path)`` — pdf is None if LibreOffice conversion
    failed. Both paths are absolute. The caller is expected to surface the PDF
    if available, else the DOCX.

    This is a side-channel artefact: it carries no signatures itself and does
    NOT mutate the enrollment row in any way (no DB writes here).
    """
    settings_logo = Path(current_app.root_path).parent.parent / "frontend" / "public" / "ccu-logo.png"
    docx_full, pdf_full = _output_paths(e)
    verify_url = public_verify_url(e.verification_code, base_url=public_base) if e.verification_code else None

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ── Header ──
    _h(doc, "Сертификат подписания", size=22)
    _h(doc, f"Договор № {e.number or '—'} от {_fmt_date(e.contract_date)}",
       size=13, color=_BRAND_CHARCOAL, bold=True)
    _p(doc, "CCU PRACTICUM · College of Caspian University",
       size=9, color=_BRAND_CHARCOAL, italic=True)
    doc.add_paragraph()

    # ── Document summary ──
    _h(doc, "Сведения о документе", size=14, color=_BRAND_CHARCOAL)
    _kv(doc, "Номер договора", e.number or "—", bold_value=True)
    _kv(doc, "Дата договора", _fmt_date(e.contract_date))
    _kv(doc, "Абитуриент", e.applicant_full_name or "—", bold_value=True)
    _kv(doc, "ИИН (скрыт)", _mask_iin(e.applicant_iin))
    _kv(doc, "Специальность", e.specialty or "—")
    if e.qualification:
        _kv(doc, "Квалификация", e.qualification)
    _kv(doc, "Сформирован сертификат", utc_now().strftime("%d.%m.%Y %H:%M UTC"))
    doc.add_paragraph()

    # ── Per-document signatures ──
    _h(doc, "Список подписей", size=14, color=_BRAND_CHARCOAL)
    sigs_by_doc: dict[str, list] = {}
    for s in e.signatures or []:
        sigs_by_doc.setdefault(s.document, []).append(s)

    # College identity — used to render the ORGANIZATION signature block explicitly
    # for every college-side signature so the reader sees the org signature, not
    # just a director's personal ЭЦП. Loaded once per certificate generation.
    college = CollegeSettings.query.first()
    college_name = (college.name_ru if college else "") or "Колледж"
    college_bin = (college.bin if college else "") or ""
    college_addr = (college.address if college else "") or ""
    college_basis = (college.director_basis if college else "") or ""

    docs_in_play = e.relevant_documents
    if not any(sigs_by_doc.get(d) for d in docs_in_play):
        _p(doc, "На момент формирования сертификата подписей нет.",
           size=10, color=_BRAND_CHARCOAL, italic=True)
    else:
        for idx, dkey in enumerate(docs_in_play):
            sigs = sigs_by_doc.get(dkey, [])
            label = DOCUMENT_LABELS.get(dkey, dkey)
            _h(doc, f"{idx + 1}. {label}", size=12, color=_BRAND_INK)
            if not sigs:
                _p(doc, "  Подписей нет.", size=10, color=_BRAND_CHARCOAL, italic=True)
                continue
            for j, s in enumerate(sorted(sigs, key=lambda x: x.created_at or 0)):
                is_college = s.signer_party == PARTY_COLLEGE
                party_label = PARTY_LABELS.get(s.signer_party, s.signer_party)
                _p(doc, f"  Подпись №{j + 1} — {party_label}", size=11, bold=True,
                   color=_BRAND_CORAL if is_college else _BRAND_INK)
                if is_college:
                    # Explicit ORGANIZATION signature block — this is what the
                    # user sees as "подпись организации", distinct from the
                    # personal (студент / родитель) signatures below.
                    _kv(doc, "    Организация", college_name, bold_value=True)
                    if college_bin:
                        _kv(doc, "    БИН", college_bin)
                    if college_addr:
                        _kv(doc, "    Адрес", college_addr)
                    _kv(doc, "    Директор", s.signer_full_name or "—", bold_value=True)
                    if college_basis:
                        _kv(doc, "    Действует на основании", college_basis)
                    _kv(doc, "    БИН подписанта (по сертификату)", _mask_iin(s.signer_iin_or_bin))
                else:
                    _kv(doc, "    Подписант", s.signer_full_name or "—", bold_value=True)
                    _kv(doc, "    ИИН (скрыт)", _mask_iin(s.signer_iin_or_bin))
                _kv(doc, "    Серийный № сертификата", _serial_tail(s.signer_serial))
                _kv(doc, "    Время подписания",
                    s.created_at.strftime("%d.%m.%Y %H:%M:%S UTC")
                    if s.created_at else "—")
                level_label = _LEVEL_RU.get(s.verification_level, s.verification_level or "—")
                _kv(doc, "    Уровень верификации", level_label)
                if s.signed_payload_sha256:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    rl = p.add_run("    SHA-256: ")
                    rl.bold = True
                    rl.font.color.rgb = _BRAND_CHARCOAL
                    rl.font.size = Pt(9)
                    rl.font.name = "Inter"
                    rv = p.add_run(s.signed_payload_sha256)
                    rv.font.size = Pt(9)
                    rv.font.color.rgb = _BRAND_INK
                    rv.font.name = "Consolas"
                doc.add_paragraph()
    doc.add_paragraph()

    # ── QR + verify URL ──
    if verify_url:
        _h(doc, "Проверка подлинности", size=14, color=_BRAND_CHARCOAL)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        qr_cell, info_cell = tbl.rows[0].cells
        qr_cell.width = Cm(4.5)
        info_cell.width = Cm(12)

        png = generate_qr_png(verify_url, box_size=10, border=2)
        qp = qr_cell.paragraphs[0]
        qp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        qp.add_run().add_picture(io.BytesIO(png), width=Cm(3.8))

        info_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        ip = info_cell.paragraphs[0]
        r = ip.add_run("Отсканируйте QR-код камерой смартфона или откройте ссылку:")
        r.font.size = Pt(10)
        r.font.color.rgb = _BRAND_CHARCOAL
        r.font.name = "Inter"

        p = info_cell.add_paragraph()
        r = p.add_run(verify_url)
        r.font.size = Pt(11)
        r.font.color.rgb = _BRAND_CORAL
        r.bold = True
        r.font.name = "Inter"

        p = info_cell.add_paragraph()
        r = p.add_run(f"Код документа: {e.verification_code}")
        r.font.size = Pt(9)
        r.font.color.rgb = _BRAND_CHARCOAL
        r.italic = True
        r.font.name = "Inter"

        p = info_cell.add_paragraph()
        r = p.add_run(
            "Страница проверки покажет всех подписантов, ФИО / ИИН и SHA-256 "
            "подписанного документа в реальном времени."
        )
        r.font.size = Pt(9)
        r.font.color.rgb = _BRAND_CHARCOAL
        r.font.name = "Inter"
    doc.add_paragraph()

    doc.save(str(docx_full))
    produced_pdf = _convert_to_pdf(docx_full)
    return docx_full, produced_pdf


def _fmt_date(value: date | None) -> str:
    return value.strftime("%d.%m.%Y") if value else "—"
