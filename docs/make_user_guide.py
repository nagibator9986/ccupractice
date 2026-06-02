"""Generate CCU PRACTICUM user guide as DOCX, then convert to PDF via LibreOffice.

Run from project root:
    python docs/make_user_guide.py

Produces:
    docs/CCU_PRACTICUM_Гайд_пользователя.docx
    docs/CCU_PRACTICUM_Гайд_пользователя.pdf
"""
from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


CORAL = RGBColor(0xE8, 0x5A, 0x3F)
CHARCOAL = RGBColor(0x46, 0x48, 0x4B)
INK = RGBColor(0x1F, 0x20, 0x24)
MUTED = RGBColor(0x6F, 0x70, 0x77)
COVER_BG = "1F2024"
CORAL_HEX = "E85A3F"
LIGHT_BG = "F7F3F1"


def _shade(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color_hex: str = "E5E7EB", size: int = 6) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _run(p, text, *, bold=False, italic=False, size=11, color=INK, name="Inter"):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = name
    # Cyrillic east-asia font fallback
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)
    if rPr.find(qn("w:rFonts")) is None:
        rPr.append(rFonts)
    return r


def heading(doc, text, *, level=1):
    sizes = {1: 22, 2: 16, 3: 13}
    spacing_before = {1: 18, 2: 14, 3: 10}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(spacing_before[level])
    p.paragraph_format.space_after = Pt(6)
    color = CORAL if level == 1 else CHARCOAL
    _run(p, text, bold=True, size=sizes[level], color=color)
    return p


def para(doc, text, *, bold=False, italic=False, size=11, color=INK, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def bullet(doc, text, *, bold_lead: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        _run(p, bold_lead, bold=True, size=11, color=INK)
        _run(p, " — " + text, size=11, color=INK)
    else:
        _run(p, text, size=11, color=INK)


def numstep(doc, n: int, title: str, body: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    _run(p, f"{n}. ", bold=True, size=12, color=CORAL)
    _run(p, title, bold=True, size=12, color=CHARCOAL)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.6)
    p2.paragraph_format.space_after = Pt(6)
    _run(p2, body, size=10.5, color=INK)


def info_box(doc, title: str, lines: list[str], *, accent: str = "E85A3F", bg: str = LIGHT_BG):
    """Render a coloured note box using a one-column table."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _shade(cell, bg)
    _set_cell_borders(cell, color_hex=accent, size=12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Title
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    _run(p, title, bold=True, size=11, color=RGBColor.from_string(accent))

    for line in lines:
        lp = cell.add_paragraph()
        lp.paragraph_format.space_after = Pt(2)
        _run(lp, line, size=10, color=INK)

    # Spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def cover_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.width = Cm(21)
    _shade(cell, COVER_BG)

    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", "2200"), ("left", "1400"), ("bottom", "2200"), ("right", "1400")):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(8)
    _run(p0, "CCU PRACTICUM", bold=True, size=10, color=CORAL)

    p1 = cell.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    _run(
        p1,
        "Гайд пользователя",
        bold=True,
        size=42,
        color=RGBColor(0xFF, 0xFF, 0xFF),
        name="Playfair Display",
    )

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(28)
    _run(
        p2,
        "Цифровое подписание трёхсторонних договоров профессиональной практики",
        size=14,
        color=RGBColor(0xCB, 0xCC, 0xCF),
    )

    p3 = cell.add_paragraph()
    _run(p3, "Для:", bold=True, size=10, color=CORAL)
    for role in (
        "Администратор колледжа",
        "Партнёр (база практики)",
        "Обучающийся (студент)",
    ):
        rp = cell.add_paragraph()
        rp.paragraph_format.space_after = Pt(2)
        _run(rp, "•  " + role, size=12, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Footer brand line
    foot = cell.add_paragraph()
    foot.paragraph_format.space_before = Pt(36)
    _run(foot, "College of Caspian University", size=10, color=RGBColor(0xCB, 0xCC, 0xCF))

    # Reset margins for the rest of the document
    doc.add_section()
    new_section = doc.sections[-1]
    new_section.top_margin = Cm(2.2)
    new_section.bottom_margin = Cm(2.2)
    new_section.left_margin = Cm(2.4)
    new_section.right_margin = Cm(2.4)


def section_intro(doc):
    heading(doc, "Что такое CCU PRACTICUM", level=1)
    para(
        doc,
        "CCU PRACTICUM — это веб-платформа Колледжа Каспийского университета для "
        "цифрового оформления и подписания трёхсторонних договоров "
        "профессиональной практики. Платформа автоматически формирует Word- и "
        "PDF-версии договора из шаблона, рассылает ссылки на подписание трём "
        "сторонам (колледж · предприятие · обучающийся), принимает электронные "
        "подписи через приложение НУЦ РК (NCALayer) и хранит подписанные документы "
        "в едином архиве.",
    )

    info_box(
        doc,
        "Три роли в системе",
        [
            "•  Администратор колледжа — заводит партнёров, студентов, формирует договоры, отправляет на подпись.",
            "•  Партнёр (предприятие) — получает ссылку, читает договор, подписывает ЭЦП.",
            "•  Обучающийся (студент) — получает ссылку, читает договор, подписывает ЭЦП.",
        ],
    )

    heading(doc, "Жизненный цикл договора", level=2)
    steps = [
        ("Черновик", "Договор создан, файл ещё не сформирован."),
        ("Сформирован", "Сгенерированы DOCX и PDF из шаблона."),
        ("Отправлен", "Созданы одноразовые ссылки на подписание для сторон."),
        ("Подписан ЭЦП", "Все три стороны подписали через NCALayer."),
        ("Скан загружен", "В систему загружен подписанный бумажный экземпляр."),
        ("Завершён", "Документооборот закрыт, договор в архиве."),
    ]
    table = doc.add_table(rows=len(steps), cols=2)
    table.autofit = False
    for i, (label, body) in enumerate(steps):
        c1 = table.cell(i, 0)
        c2 = table.cell(i, 1)
        c1.width = Cm(4.5)
        c2.width = Cm(11.5)
        _shade(c1, LIGHT_BG)
        _set_cell_borders(c1, "E5E7EB")
        _set_cell_borders(c2, "E5E7EB")
        p1 = c1.paragraphs[0]
        _run(p1, f"{i+1}. {label}", bold=True, size=10.5, color=CORAL)
        p2 = c2.paragraphs[0]
        _run(p2, body, size=10.5, color=INK)
    doc.add_paragraph()


def section_admin(doc):
    doc.add_page_break()
    heading(doc, "Для администратора колледжа", level=1)
    para(
        doc,
        "Администратор управляет реестрами и запускает документооборот. Доступ "
        "к разделу «Настройки» и операциям записи (создание, редактирование, "
        "удаление, генерация, отправка на подписание) есть только у роли "
        "Администратор. Роль «Просмотр» (viewer) видит данные, но не может "
        "ничего изменить.",
    )

    info_box(
        doc,
        "Вход в систему",
        [
            "Откройте адрес платформы и нажмите «Войти».",
            "По умолчанию: admin@ccu.kz / admin123 (смените пароль при первом входе через раздел «Настройки»).",
        ],
    )

    heading(doc, "Шаги работы", level=2)

    numstep(doc, 1,
        "Заполните реквизиты колледжа",
        "Меню → Настройки. Укажите наименование, БИН, банковские реквизиты, ФИО "
        "директора, адрес. Эти данные будут автоматически подставлены в каждый "
        "формируемый договор.")

    numstep(doc, 2,
        "Добавьте партнёра (базу практики)",
        "Меню → Партнёры → «Добавить партнёра». Введите наименование организации, "
        "БИН, юридический и фактический адрес, ФИО руководителя, должность, "
        "контактный email и телефон, банковские реквизиты. Количество мест и "
        "статус договора помогают фильтровать партнёров.")

    numstep(doc, 3,
        "Добавьте студента",
        "Меню → Студенты → «Добавить студента». Заполните ФИО, ИИН, группу, "
        "специальность и её код, период практики, форму обучения, вид практики, "
        "руководителей от колледжа и от предприятия. При наличии — данные "
        "законного представителя.")

    numstep(doc, 4,
        "Создайте договор",
        "Меню → Договоры → «+ Новый договор». Выберите партнёра и студента, "
        "укажите дату. Номер ПП-ГОД-NNN присваивается автоматически.")

    numstep(doc, 5,
        "Сформируйте файлы",
        "Откройте карточку договора и нажмите «Сформировать договор». Платформа "
        "сгенерирует DOCX из шаблона и автоматически сконвертирует в PDF.")

    numstep(doc, 6,
        "Отправьте на подписание",
        "В разделе «Документооборот» нажмите «Отправить партнёру и студенту» "
        "или «Отправить всем». Будут созданы три одноразовые ссылки. Скопируйте "
        "их из модального окна и передайте получателям (email / WhatsApp / "
        "Telegram — кнопки уже встроены).")

    numstep(doc, 7,
        "Подпишите от имени колледжа",
        "Рядом со строкой «Колледж» нажмите «Подписать здесь». Откроется "
        "NCALayer — выберите ключ ЭЦП директора и введите PIN. Подпись будет "
        "привязана к договору.")

    numstep(doc, 8,
        "Контролируйте процесс",
        "В карточке договора видны статусы каждой стороны: pending / viewed / "
        "signed / revoked. Можно «Перевыпустить» (новый токен) или «Отозвать» "
        "просроченную или ошибочно отправленную ссылку.")

    numstep(doc, 9,
        "Загрузите скан и закройте договор",
        "Если параллельно подписали бумажную версию — загрузите её скан в "
        "разделе «Управление статусом». Когда документооборот окончен — "
        "переведите статус в «Завершён». Договор останется в Электронном архиве.")

    info_box(
        doc,
        "Где искать готовые документы",
        [
            "Меню → Электронный архив (списком или деревом по годам и партнёрам).",
            "Скачать DOCX, PDF и Отчёт о подписях можно из карточки любого договора.",
        ],
    )

    heading(doc, "Шаблон договора", level=2)
    para(
        doc,
        "Шаблон по умолчанию построен из исходного «Договора 3-х сторон» и "
        "поддерживает переменные Jinja2. В разделе «Настройки» вы можете загрузить "
        "собственный .docx-шаблон. Доступные переменные:",
    )
    for tpl in (
        "{{ contract.number }}, {{ contract.date }}, {{ contract.practice_start }}, {{ contract.practice_end }}",
        "{{ college.name_ru }}, {{ college.address }}, {{ college.bin }}, {{ college.director_full_name }}",
        "{{ partner.organization_name }}, {{ partner.bin }}, {{ partner.director_position }}, {{ partner.director_full_name }}",
        "{{ student.full_name }}, {{ student.iin }}, {{ student.group_name }}, {{ student.specialty }}",
    ):
        bullet(doc, tpl)


def section_partner(doc):
    doc.add_page_break()
    heading(doc, "Для партнёра (предприятия)", level=1)
    para(
        doc,
        "Вам не нужен логин и пароль. Колледж отправляет одноразовую ссылку, "
        "по которой можно посмотреть договор и подписать его электронной "
        "подписью (ЭЦП), используя приложение НУЦ РК (NCALayer).",
    )

    heading(doc, "Что нужно подготовить", level=2)
    bullet(doc, "Электронную подпись (ЭЦП) руководителя предприятия — файл PKCS12, ID-карта или JaCarta-токен.")
    bullet(doc, "Установленное и запущенное приложение NCALayer — https://pki.gov.kz/ncalayer/.")
    bullet(doc, "Современный браузер (Chrome, Edge, Safari, Firefox последних версий).")

    heading(doc, "Шаги подписания", level=2)

    numstep(doc, 1,
        "Откройте ссылку из письма",
        "Колледж пришлёт письмо с уникальной ссылкой вида .../sign/<токен>. "
        "Откройте её — авторизация не требуется.")

    numstep(doc, 2,
        "Прочитайте договор",
        "На открывшейся странице вы увидите все реквизиты договора: номер, "
        "дату, информацию о студенте, периоде практики и о вашей организации. "
        "Скачайте полный текст в DOCX или PDF, чтобы внимательно ознакомиться.")

    numstep(doc, 3,
        "Запустите NCALayer",
        "Убедитесь, что приложение NCALayer запущено (иконка в трее/панели). "
        "Платформа автоматически проверит его доступность и покажет зелёный "
        "индикатор «NCALayer обнаружен».")

    numstep(doc, 4,
        "Нажмите «Подписать ЭЦП»",
        "Откроется окно NCALayer. Выберите хранилище ключа (PKCS12 / "
        "удостоверение / токен), укажите путь к файлу подписи, введите PIN-код. "
        "После подтверждения подпись будет автоматически отправлена в систему "
        "колледжа.")

    numstep(doc, 5,
        "Готово",
        "После успешной подписи появится зелёное уведомление «Договор подписан "
        "с вашей стороны». Ссылка остаётся активной для просмотра договора и "
        "его статуса, но повторно подписать по ней уже нельзя.")

    info_box(
        doc,
        "Безопасность ссылки",
        [
            "Ссылка одноразовая и привязана только к вашему предприятию.",
            "Действует до даты, указанной в правой панели (по умолчанию 30 дней).",
            "Сертификат вашего ЭЦП фиксируется в системе, а в архиве хранится SHA-256 хэш подписанного файла.",
            "При утере ссылки — попросите колледж нажать «Перевыпустить».",
        ],
    )


def section_student(doc):
    doc.add_page_break()
    heading(doc, "Для обучающегося (студента)", level=1)
    para(
        doc,
        "Логин не нужен. Вы получите от колледжа короткую ссылку — по ней "
        "можно прочитать свой договор о профессиональной практике и подписать "
        "его собственной ЭЦП через NCALayer.",
    )

    heading(doc, "Что нужно подготовить", level=2)
    bullet(doc, "Личную ЭЦП — обычно это файл RSA*.p12, выданный НУЦ РК или полученный через eGov.kz.")
    bullet(doc, "Установленное приложение NCALayer (https://pki.gov.kz/ncalayer/), запущенное на компьютере.")
    bullet(doc, "Современный браузер (Chrome, Edge, Safari, Firefox).")

    heading(doc, "Шаги подписания", level=2)

    numstep(doc, 1,
        "Откройте ссылку из сообщения",
        "Ссылка вида .../sign/<токен> приходит от колледжа в почте, WhatsApp "
        "или Telegram. Откройте её в браузере на ноутбуке/ПК (NCALayer "
        "работает только на десктопе).")

    numstep(doc, 2,
        "Проверьте свои данные",
        "Убедитесь, что ваши ФИО, ИИН, группа, специальность, период практики и "
        "название предприятия совпадают с реальностью. Если есть ошибка — "
        "не подписывайте и сообщите в колледж, чтобы они исправили данные.")

    numstep(doc, 3,
        "Запустите NCALayer",
        "Иконка приложения должна быть в трее (рядом с часами). На странице "
        "договора в блоке «Электронная подпись» убедитесь, что отображается "
        "«NCALayer обнаружен на вашем устройстве».")

    numstep(doc, 4,
        "Подпишите ЭЦП",
        "Нажмите большую кнопку «Подписать ЭЦП». В окне NCALayer выберите ваш "
        "сертификат (обычно файл RSA*.p12), введите пароль ЭЦП. После "
        "подтверждения подпись передаётся в систему.")

    numstep(doc, 5,
        "Подтверждение",
        "Появится зелёный блок «Договор подписан с вашей стороны». Ссылка "
        "останется активной для просмотра, но повторно подписать вы не сможете. "
        "Когда все три стороны (колледж, предприятие и вы) подпишут — договор "
        "автоматически перейдёт в статус «Подписан».")

    info_box(
        doc,
        "Если у вас нет ЭЦП",
        [
            "Получите её бесплатно через портал eGov.kz («Получить онлайн-ЭЦП») или в ЦОНе.",
            "Несовершеннолетним: подписать может законный представитель (укажите его в данных).",
            "Сообщите в колледж — администратор подскажет, нужна ли вам ЭЦП конкретно для этого договора.",
        ],
    )


def section_faq(doc):
    doc.add_page_break()
    heading(doc, "Часто задаваемые вопросы", level=1)

    faqs = [
        ("Безопасно ли подписывать через NCALayer в браузере?",
         "Да. NCALayer — официальное приложение НУЦ РК. Подпись CMS/CAdES "
         "формируется локально на вашем компьютере, приватный ключ никогда "
         "не покидает ваш токен или файл .p12. Платформа получает только "
         "криптографическую структуру подписи и сертификат."),
        ("Что если NCALayer не определяется?",
         "1) Убедитесь, что приложение запущено и его иконка в трее.\n"
         "2) Откройте https://127.0.0.1:13579 — если браузер ругается на "
         "сертификат, примите его (один раз).\n"
         "3) Проверьте, что NCALayer обновлён до последней версии."),
        ("Можно ли подписать с телефона?",
         "Технически NCALayer есть для мобильных устройств, но рекомендуется "
         "подписывать с ноутбука/ПК — там удобнее работать с ЭЦП и файлами."),
        ("Я случайно подписал не то — что делать?",
         "Срочно сообщите администратору колледжа. Он отзовёт подпись, исправит "
         "данные и выпустит новую ссылку на подписание (старый токен станет "
         "недействителен)."),
        ("Где хранятся подписанные договоры?",
         "В электронном архиве платформы под структурой: "
         "«Профессиональная практика → Год → Партнёр → Договоры». Доступ — "
         "только у Администратора. Каждый договор сопровождается DOCX, PDF и "
         "автоматически генерируемым Отчётом о подписях, в котором фиксируются "
         "ФИО, ИИН/БИН подписантов, серийные номера их сертификатов и SHA-256 "
         "хэш подписанного файла."),
        ("Сколько действует ссылка?",
         "По умолчанию 30 дней с момента создания. После истечения откроется "
         "сообщение «Срок действия ссылки истёк», и потребуется попросить "
         "у колледжа новую."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        _run(p, q, bold=True, size=11.5, color=CHARCOAL)
        for line in a.split("\n"):
            ap = doc.add_paragraph()
            ap.paragraph_format.left_indent = Cm(0.3)
            ap.paragraph_format.space_after = Pt(2)
            _run(ap, line, size=10.5, color=INK)

    heading(doc, "Контакты", level=2)
    para(
        doc,
        "По вопросам работы платформы — пишите администратору колледжа, чья "
        "контактная информация указана в шапке вашего письма со ссылкой, либо "
        "в разделе «Настройки → Реквизиты колледжа» внутри платформы.",
    )

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(20)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(foot, "© College of Caspian University · CCU PRACTICUM", size=9, color=MUTED)


def build_docx(path: Path) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Inter"
    style.font.size = Pt(11)

    cover_page(doc)
    section_intro(doc)
    section_admin(doc)
    section_partner(doc)
    section_student(doc)
    section_faq(doc)

    doc.save(path)
    return path


def convert_to_pdf(docx_path: Path) -> Path | None:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        print("soffice not found — PDF not generated.")
        return None
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
        check=True,
    )
    return pdf_path if pdf_path.exists() else None


def main():
    out_dir = Path(__file__).resolve().parent
    docx_path = out_dir / "CCU_PRACTICUM_Гайд_пользователя.docx"
    build_docx(docx_path)
    pdf = convert_to_pdf(docx_path)
    print(f"DOCX: {docx_path}")
    if pdf:
        print(f"PDF:  {pdf}")


if __name__ == "__main__":
    main()
